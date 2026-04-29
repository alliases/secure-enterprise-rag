# File: app/ingestion/parser.py
# Purpose: Extract raw text and metadata from uploaded PDF and DOCX files.

from dataclasses import dataclass
from pathlib import Path
from typing import Any, cast

import docx  # type: ignore[import-untyped]
import fitz  # PyMuPDF # type: ignore[import-untyped]

from app.logging_config.setup import get_logger

logger = get_logger(__name__)


class ParseError(Exception):
    """Custom exception raised when document parsing fails."""

    pass


@dataclass
class ParsedDocument:
    """Data container for the extracted text and its associated metadata."""

    text: str
    metadata: dict[str, Any]


def parse_pdf(file_path: Path) -> tuple[str, int]:
    """
    Extracts text from a PDF file using PyMuPDF.
    Returns a tuple containing the concatenated text and the total page count.
    """
    text_blocks: list[str] = []
    page_count = 0

    try:
        # Context manager ensures the document stream is closed properly
        with fitz.open(str(file_path)) as doc:
            page_count = len(doc)
            for page in doc:
                # Explicitly cast to str to resolve PyMuPDF Union return type
                text = cast(str, page.get_text())  # type: ignore[reportUnknownMemberType]
                text_blocks.append(text)
    except Exception as e:
        logger.error("PDF parsing failed", file_path=str(file_path), error=str(e))
        raise ParseError(f"Failed to parse PDF document: {e!s}") from e

    return "\n\n".join(text_blocks), page_count


def parse_docx(file_path: Path) -> str:
    """
    Extracts text from a DOCX file using python-docx.
    Iterates through paragraphs and concatenates non-empty text blocks.
    """
    text_blocks: list[str] = []

    try:
        # Ignore strict typing errors for completely untyped python-docx library
        doc = docx.Document(str(file_path))  # type: ignore

        # Extract text from standard paragraphs
        for para in doc.paragraphs:  # type: ignore
            text = str(para.text).strip()  # type: ignore
            if text:
                text_blocks.append(text)

        # Extract text from tables to prevent data loss in structured documents
        for table in doc.tables:  # type: ignore
            for row in table.rows:  # type: ignore
                row_data: list[str] = [
                    str(cell.text).strip()  # type: ignore
                    for cell in row.cells  # type: ignore
                    if str(cell.text).strip()  # type: ignore
                ]
                if row_data:
                    text_blocks.append(" | ".join(row_data))

    except Exception as e:
        logger.error("DOCX parsing failed", file_path=str(file_path), error=str(e))
        raise ParseError(f"Failed to parse DOCX document: {e!s}") from e

    return "\n\n".join(text_blocks)


def parse_document(file_path: Path, file_name: str, file_type: str) -> ParsedDocument:
    """
    Main entry point for document parsing.
    Delegates to specific parsers based on the provided file_type.
    """
    if not file_path.exists():
        raise ParseError(f"File not found: {file_path}")

    file_type_lower = file_type.lower()
    extracted_text = ""
    page_count = 0

    if file_type_lower == "pdf":
        extracted_text, page_count = parse_pdf(file_path)
    elif file_type_lower in ["docx", "doc"]:
        extracted_text = parse_docx(file_path)
        # DOCX lacks a strict pagination concept without rendering engines
        page_count = 1
    else:
        logger.warning("Unsupported file type attempted", file_type=file_type)
        raise ParseError(f"Unsupported file type: {file_type}")

    metadata = {
        "file_name": file_name,
        "file_type": file_type_lower,
        "page_count": page_count,
    }

    return ParsedDocument(text=extracted_text, metadata=metadata)
